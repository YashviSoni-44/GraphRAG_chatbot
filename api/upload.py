# from flask import Blueprint, request, jsonify, current_app
# import os
# import pdfplumber
# import pandas as pd
# import docx
# from neo4j import GraphDatabase, basic_auth
# from dotenv import load_dotenv
# import nltk
# from nltk.tokenize import sent_tokenize
# import logging

# # Use bundled Punkt tokenizer data, do NOT download at runtime
# nltk.data.path.append(os.path.join(os.path.dirname(__file__), "..", "nltk_data"))

# load_dotenv()

# upload_bp = Blueprint("upload", __name__)

# # Use /tmp/uploads or configure your persistent upload folder
# BASE_UPLOAD_FOLDER = "/tmp/uploads"
# os.makedirs(BASE_UPLOAD_FOLDER, exist_ok=True)

# NEO4J_URI = os.getenv("NEO4J_URI")
# NEO4J_USER = os.getenv("NEO4J_USER")
# NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")

# ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx"}

# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text(file_path, filename):
#     ext = filename.split(".")[-1].lower()
#     try:
#         if ext == "pdf":
#             with pdfplumber.open(file_path) as pdf:
#                 texts = [page.extract_text() for page in pdf.pages if page.extract_text()]
#                 text = "\n".join(texts)
#         elif ext == "docx":
#             doc = docx.Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#         elif ext == "txt":
#             with open(file_path, "r", encoding="utf-8") as f:
#                 text = f.read()
#         elif ext == "csv":
#             df = pd.read_csv(file_path)
#             text = df.to_string(index=False)
#         elif ext == "xlsx":
#             df = pd.read_excel(file_path)
#             text = df.to_string(index=False)
#         else:
#             text = ""
#     except Exception as e:
#         current_app.logger.error(f"Error extracting text from {filename}: {e}")
#         return "", f"extractor_error: {str(e)}"
#     return text, None

# def create_graph_for_file(session, session_id, filename, text):
#     # Compose a unique Neo4j node file name combining session ID to isolate users
#     node_file_name = f"{session_id}__{filename}"

#     sentences = sent_tokenize(text)
#     if not sentences:
#         return 0
#     session.run("MERGE (f:File {name: $filename, sessionId: $session_id})",
#                 {"filename": node_file_name, "session_id": session_id})

#     # ChatSession likewise scoped by sessionId and filename
#     session.run("""
#         MERGE (cs:ChatSession {file_name: $filename, sessionId: $session_id})
#         WITH cs
#         MATCH (f:File {name: $filename, sessionId: $session_id})
#         MERGE (f)-[:HAS_CHAT_SESSION]->(cs)
#     """, {"filename": node_file_name, "session_id": session_id})

#     for idx, sentence in enumerate(sentences):
#         sentence_id = f"{node_file_name}_{idx}"
#         session.run("""
#             MERGE (s:Sentence {id: $id, sessionId: $session_id})
#             ON CREATE SET s.text = $text
#         """, {"id": sentence_id, "text": sentence, "session_id": session_id})

#         session.run("""
#             MATCH (f:File {name: $filename, sessionId: $session_id}), 
#                   (s:Sentence {id: $id, sessionId: $session_id})
#             MERGE (f)-[:CONTAINS]->(s)
#         """, {"filename": node_file_name, "id": sentence_id, "session_id": session_id})

#         if idx > 0:
#             prev_id = f"{node_file_name}_{idx - 1}"
#             session.run("""
#                 MATCH (a:Sentence {id: $prev_id, sessionId: $session_id}), 
#                       (b:Sentence {id: $curr_id, sessionId: $session_id})
#                 MERGE (a)-[:NEXT]->(b)
#             """, {"prev_id": prev_id, "curr_id": sentence_id, "session_id": session_id})

#     return len(sentences)

# @upload_bp.route("", methods=["POST"])
# def upload_file():
#     # Get session ID from header or form data (prefer header)
#     session_id = request.headers.get("X-Session-Id") or request.form.get("sessionId")
#     if not session_id:
#         return jsonify({"error": "Missing session ID"}), 400

#     # Accept files under 'file[]' or 'file' keys for flexibility
#     files = request.files.getlist('file[]') or request.files.getlist('file') or []
#     if not files:
#         current_app.logger.warning("Upload attempt with no files received")
#         return jsonify({"error": "No files uploaded"}), 400

#     uploaded_files = []
#     errors = []

#     driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
#     with driver.session() as session:
#         for file in files:
#             if not file or file.filename == '':
#                 errors.append("One or more files missing filename.")
#                 continue

#             current_app.logger.info(f"Processing uploaded file: {file.filename}")

#             rel_path = file.filename.replace("\\", "/")  # Normalize slashes

#             if not allowed_file(rel_path):
#                 errors.append(f"{rel_path}: invalid file type")
#                 continue

#             # Save inside a subfolder named by session ID for isolation
#             save_dir = os.path.join(BASE_UPLOAD_FOLDER, session_id)
#             os.makedirs(save_dir, exist_ok=True)
#             save_path = os.path.join(save_dir, rel_path)

#             # Ensure parent dirs exist
#             parent_dir = os.path.dirname(save_path)
#             if not os.path.exists(parent_dir):
#                 os.makedirs(parent_dir, exist_ok=True)

#             try:
#                 file.save(save_path)
#             except Exception as e:
#                 errors.append(f"{rel_path}: failed to save ({e})")
#                 continue

#             text, err = extract_text(save_path, rel_path)
#             if err:
#                 errors.append(f"{rel_path}: {err}")
#                 continue

#             if not text.strip():
#                 errors.append(f"{rel_path}: empty or unreadable")
#                 continue

#             try:
#                 count = create_graph_for_file(session, session_id, rel_path, text)
#                 if count == 0:
#                     errors.append(f"{rel_path}: no sentences could be parsed from extracted text")
#                     continue
#                 uploaded_files.append(rel_path)
#             except Exception as e:
#                 errors.append(f"{rel_path}: graph creation error: {str(e)}")

#     driver.close()

#     if not uploaded_files:
#         return jsonify({"error": "No valid files were uploaded.", "errors": errors}), 400

#     response = {
#         "message": f"Files uploaded ({', '.join(uploaded_files)}) and graph updated.",
#         "success": uploaded_files
#     }
#     if errors:
#         response["errors"] = errors
#     return jsonify(response), 200

# @upload_bp.route("/files", methods=["GET"])
# def list_uploaded_files():
#     session_id = request.headers.get("X-Session-Id") or request.args.get("sessionId")
#     if not session_id:
#         return jsonify({"error": "Missing session ID"}), 400

#     driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
#     files = []
#     try:
#         with driver.session() as session:
#             result = session.run("""
#                 MATCH (f:File {sessionId: $session_id}) 
#                 RETURN f.name AS name ORDER BY name
#             """, {"session_id": session_id})
#             # Remove sessionId prefix from filenames before returning
#             files = [record["name"].split("__", 1)[1] if "__" in record["name"] else record["name"] for record in result]
#     except Exception as e:
#         current_app.logger.error(f"Failed to fetch files list: {e}")
#         return jsonify({"error": "Failed to fetch files list"}), 500
#     finally:
#         driver.close()
#     return jsonify({"files": files})

# @upload_bp.route("/delete", methods=["POST"])
# def delete_file_and_chat():
#     data = request.get_json()
#     filename = data.get("file_name")
#     session_id = request.headers.get("X-Session-Id") or data.get("sessionId")
#     if not filename or not session_id:
#         return jsonify({"error": "file_name and sessionId are required for deletion"}), 400

#     # Compose session-scoped filename key as stored in Neo4j
#     node_file_name = f"{session_id}__{filename}"

#     driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
#     errors = []
#     try:
#         with driver.session() as session:
#             # Delete all sentences belonging to file and session
#             session.run("""
#                 MATCH (f:File {name: $filename, sessionId: $session_id})-[:CONTAINS]->(s:Sentence {sessionId: $session_id})
#                 DETACH DELETE s
#             """, {"filename": node_file_name, "session_id": session_id})

#             # Delete chat session(s) and messages scoped by session and filename
#             session.run("""
#                 MATCH (cs:ChatSession {file_name: $filename, sessionId: $session_id})-[:HAS_MESSAGE]->(m:ChatMessage)
#                 DETACH DELETE m
#             """, {"filename": node_file_name, "session_id": session_id})
#             session.run("""
#                 MATCH (cs:ChatSession {file_name: $filename, sessionId: $session_id})
#                 DETACH DELETE cs
#             """, {"filename": node_file_name, "session_id": session_id})

#             # Delete file node itself
#             session.run("""
#                 MATCH (f:File {name: $filename, sessionId: $session_id})
#                 DETACH DELETE f
#             """, {"filename": node_file_name, "session_id": session_id})

#         # Remove uploaded file itself
#         file_path = os.path.join(BASE_UPLOAD_FOLDER, session_id, filename)
#         if os.path.exists(file_path):
#             try:
#                 os.remove(file_path)
#                 # Clean up empty dirs upwards
#                 dir_path = os.path.dirname(file_path)
#                 while dir_path and dir_path != os.path.join(BASE_UPLOAD_FOLDER, session_id) and os.path.isdir(dir_path):
#                     if not os.listdir(dir_path):
#                         os.rmdir(dir_path)
#                         dir_path = os.path.dirname(dir_path)
#                     else:
#                         break
#             except Exception as e:
#                 errors.append(f"File system error: {e}")

#         driver.close()
#         msg = f"Deleted file and chat for '{filename}'."
#         if errors:
#             msg += " Some issues: " + "; ".join(errors)
#         return jsonify({"message": msg})
#     except Exception as e:
#         driver.close()
#         return jsonify({"error": "Failed to delete file/chat: " + str(e)}), 500


#------------------------------------------------------------------------------------------------------



from flask import Blueprint, request, jsonify, current_app
import os
import pdfplumber
import pandas as pd
import docx
from neo4j import GraphDatabase, basic_auth
from dotenv import load_dotenv
import nltk
from nltk.tokenize import sent_tokenize
import logging

# Use bundled Punkt tokenizer data, do NOT download at runtime
nltk.data.path.append(os.path.join(os.path.dirname(__file__), "..", "nltk_data"))

load_dotenv()

upload_bp = Blueprint("upload", __name__)

# Use /tmp/uploads or configure your persistent upload folder
BASE_UPLOAD_FOLDER = "/tmp/uploads"
os.makedirs(BASE_UPLOAD_FOLDER, exist_ok=True)

NEO4J_URI = os.getenv("NEO4J_URI")
NEO4J_USER = os.getenv("NEO4J_USER")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx"}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text(file_path, filename):
    ext = filename.split(".")[-1].lower()
    try:
        if ext == "pdf":
            with pdfplumber.open(file_path) as pdf:
                texts = [page.extract_text() for page in pdf.pages if page.extract_text()]
                text = "\n".join(texts)
        elif ext == "docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif ext == "csv":
            df = pd.read_csv(file_path)
            text = df.to_string(index=False)
        elif ext == "xlsx":
            df = pd.read_excel(file_path)
            text = df.to_string(index=False)
        else:
            text = ""
    except Exception as e:
        current_app.logger.error(f"Error extracting text from {filename}: {e}")
        return "", f"extractor_error: {str(e)}"
    return text, None

def create_graph_for_file(session, session_id, filename, text):
    node_file_name = f"{session_id}__{filename}"
    sentences = sent_tokenize(text)
    if not sentences:
        return 0
    session.run("MERGE (f:File {name: $filename, sessionId: $session_id})",
                {"filename": node_file_name, "session_id": session_id})

    session.run("""
        MERGE (cs:ChatSession {file_name: $filename, sessionId: $session_id})
        WITH cs
        MATCH (f:File {name: $filename, sessionId: $session_id})
        MERGE (f)-[:HAS_CHAT_SESSION]->(cs)
    """, {"filename": node_file_name, "session_id": session_id})

    for idx, sentence in enumerate(sentences):
        sentence_id = f"{node_file_name}_{idx}"
        session.run("""
            MERGE (s:Sentence {id: $id, sessionId: $session_id})
            ON CREATE SET s.text = $text
        """, {"id": sentence_id, "text": sentence, "session_id": session_id})

        session.run("""
            MATCH (f:File {name: $filename, sessionId: $session_id}), 
                  (s:Sentence {id: $id, sessionId: $session_id})
            MERGE (f)-[:CONTAINS]->(s)
        """, {"filename": node_file_name, "id": sentence_id, "session_id": session_id})

        if idx > 0:
            prev_id = f"{node_file_name}_{idx - 1}"
            session.run("""
                MATCH (a:Sentence {id: $prev_id, sessionId: $session_id}), 
                      (b:Sentence {id: $curr_id, sessionId: $session_id})
                MERGE (a)-[:NEXT]->(b)
            """, {"prev_id": prev_id, "curr_id": sentence_id, "session_id": session_id})

    return len(sentences)

@upload_bp.route("", methods=["POST"])
def upload_file():
    session_id = request.headers.get("X-Session-Id") or request.form.get("sessionId")
    if not session_id:
        return jsonify({"error": "Missing session ID"}), 400

    files = request.files.getlist('file[]') or request.files.getlist('file') or []
    if not files:
        current_app.logger.warning("Upload attempt with no files received")
        return jsonify({"error": "No files uploaded"}), 400

    uploaded_files = []
    errors = []

    driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
    with driver.session() as session:
        for file in files:
            if not file or file.filename == '':
                errors.append("One or more files missing filename.")
                continue

            current_app.logger.info(f"Processing uploaded file: {file.filename}")

            rel_path = file.filename.replace("\\", "/")  # Normalize slashes

            if not allowed_file(rel_path):
                errors.append(f"{rel_path}: invalid file type")
                continue

            save_dir = os.path.join(BASE_UPLOAD_FOLDER, session_id)
            os.makedirs(save_dir, exist_ok=True)
            save_path = os.path.join(save_dir, rel_path)

            parent_dir = os.path.dirname(save_path)
            if not os.path.exists(parent_dir):
                os.makedirs(parent_dir, exist_ok=True)

            try:
                file.save(save_path)
            except Exception as e:
                errors.append(f"{rel_path}: failed to save ({e})")
                continue

            text, err = extract_text(save_path, rel_path)
            if err:
                errors.append(f"{rel_path}: {err}")
                continue

            if not text.strip():
                errors.append(f"{rel_path}: empty or unreadable")
                continue

            try:
                count = create_graph_for_file(session, session_id, rel_path, text)
                if count == 0:
                    errors.append(f"{rel_path}: no sentences could be parsed from extracted text")
                    continue
                uploaded_files.append(rel_path)
            except Exception as e:
                errors.append(f"{rel_path}: graph creation error: {str(e)}")

    driver.close()

    if not uploaded_files:
        return jsonify({"error": "No valid files were uploaded.", "errors": errors}), 400

    response = {
        "message": f"Files uploaded ({', '.join(uploaded_files)}) and graph updated.",
        "success": uploaded_files
    }
    if errors:
        response["errors"] = errors
    return jsonify(response), 200

@upload_bp.route("/files", methods=["GET"])
def list_uploaded_files():
    session_id = request.headers.get("X-Session-Id") or request.args.get("sessionId")
    if not session_id:
        return jsonify({"error": "Missing session ID"}), 400

    driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
    files = []
    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (f:File {sessionId: $session_id}) 
                RETURN f.name AS name ORDER BY name
            """, {"session_id": session_id})
            files = [record["name"].split("__", 1)[1] if "__" in record["name"] else record["name"] for record in result]
    except Exception as e:
        current_app.logger.error(f"Failed to fetch files list: {e}")
        return jsonify({"error": "Failed to fetch files list"}), 500
    finally:
        driver.close()
    return jsonify({"files": files})

@upload_bp.route("/delete", methods=["POST"])
def delete_file_and_chat():
    data = request.get_json()
    filename = data.get("file_name")
    session_id = request.headers.get("X-Session-Id") or data.get("sessionId")
    if not filename or not session_id:
        return jsonify({"error": "file_name and sessionId are required for deletion"}), 400

    node_file_name = f"{session_id}__{filename}"

    driver = GraphDatabase.driver(NEO4J_URI, auth=basic_auth(NEO4J_USER, NEO4J_PASSWORD))
    errors = []
    try:
        with driver.session() as session:
            session.run("""
                MATCH (f:File {name: $filename, sessionId: $session_id})-[:CONTAINS]->(s:Sentence {sessionId: $session_id})
                DETACH DELETE s
            """, {"filename": node_file_name, "session_id": session_id})

            session.run("""
                MATCH (cs:ChatSession {file_name: $filename, sessionId: $session_id})-[:HAS_MESSAGE]->(m:ChatMessage)
                DETACH DELETE m
            """, {"filename": node_file_name, "session_id": session_id})
            session.run("""
                MATCH (cs:ChatSession {file_name: $filename, sessionId: $session_id})
                DETACH DELETE cs
            """, {"filename": node_file_name, "session_id": session_id})

            session.run("""
                MATCH (f:File {name: $filename, sessionId: $session_id})
                DETACH DELETE f
            """, {"filename": node_file_name, "session_id": session_id})

        file_path = os.path.join(BASE_UPLOAD_FOLDER, session_id, filename)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                dir_path = os.path.dirname(file_path)
                while dir_path and dir_path != os.path.join(BASE_UPLOAD_FOLDER, session_id) and os.path.isdir(dir_path):
                    if not os.listdir(dir_path):
                        os.rmdir(dir_path)
                        dir_path = os.path.dirname(dir_path)
                    else:
                        break
            except Exception as e:
                errors.append(f"File system error: {e}")

        driver.close()
        msg = f"Deleted file and chat for '{filename}'."
        if errors:
            msg += " Some issues: " + "; ".join(errors)
        return jsonify({"message": msg})
    except Exception as e:
        driver.close()
        return jsonify({"error": "Failed to delete file/chat: " + str(e)}), 500
